import os
import requests
import random
import hashlib
import time
import re
from bs4 import BeautifulSoup
from urllib.parse import urljoin, urlparse
from PyPDF2 import PdfReader
import docx
from collections import deque

# Constants
EXCLUDE_EXTENSIONS = (
    ".jpg", ".jpeg", ".png", ".gif", ".mp4", ".avi", ".mov", ".wmv",
    ".mp3", ".wav", ".ogg", ".webm", ".svg", ".ico", ".css", ".js",
    ".ttf", ".woff", ".woff2", ".eot", ".zip", ".tar", ".gz", ".rar",
)

TEXT_EXTENSIONS = (".pdf", ".doc", ".docx", ".txt")

# Common patterns to identify and remove non-essential content
COMMON_NOISE_PATTERNS = [
    "cookie", "privacy policy", "terms of use", "all rights reserved",
    "©", "copyright", "follow us", "subscribe", "newsletter",
    "share this", "comments", "related posts", "tags:", "categories:"
]

# Social media domains to exclude from scraping
SOCIAL_MEDIA_DOMAINS = [
    "facebook.com", "twitter.com", "instagram.com", "linkedin.com",
    "youtube.com", "pinterest.com", "tiktok.com", "reddit.com",
    "tumblr.com", "snapchat.com", "whatsapp.com", "telegram.org"
]

def is_valid_url(url, base_url):
    """Check if the URL is valid for scraping."""
    parsed_url = urlparse(url)
    parsed_base = urlparse(base_url)
    
    # Check if it's an external link
    if parsed_url.netloc and parsed_url.netloc != parsed_base.netloc:
        return False
    
    # Check if it's a social media link
    if any(social_domain in parsed_url.netloc for social_domain in SOCIAL_MEDIA_DOMAINS):
        return False
    
    # Check if it has excluded extensions
    if url.lower().endswith(EXCLUDE_EXTENSIONS):
        return False
    
    # Skip mailto: links and javascript: links
    if url.startswith(("mailto:", "javascript:", "tel:")):
        return False
    
    return True

def compute_hash(content):
    """Compute SHA-256 hash of content."""
    return hashlib.sha256(content.encode('utf-8')).hexdigest()

def clean_text(text):
    """Clean the extracted text by removing noise and normalizing whitespace."""
    # Remove common noise patterns
    for pattern in COMMON_NOISE_PATTERNS:
        text = text.replace(pattern, "")
    
    # Normalize whitespace
    # Replace multiple spaces with a single space
    text = re.sub(r' +', ' ', text)
    
    # Remove spaces at the beginning of lines
    text = re.sub(r'\n +', '\n', text)
    
    # Normalize paragraph breaks (ensure exactly one blank line between paragraphs)
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    # Remove trailing whitespace at the end of lines
    text = re.sub(r' +\n', '\n', text)
    
    return text.strip()

def extract_main_content(soup):
    """Extract the main content from the page with proper formatting."""
    # First attempt: Look for main content blocks
    main_content_element = None
    
    # Try to find the main article content
    article = soup.find('article')
    if article:
        main_content_element = article
    
    # If no article found, try common content containers
    if not main_content_element:
        content_containers = soup.find_all(['div', 'section'], class_=lambda c: c and any(
            content_class in c.lower() for content_class in 
            ['content', 'article', 'blog', 'post', 'entry', 'main', 'body']
        ))
        
        largest_text_length = 0
        for container in content_containers:
            container_text = container.get_text(strip=True)
            if len(container_text) > largest_text_length:
                largest_text_length = len(container_text)
                main_content_element = container
    
    # Format the content with proper structure
    formatted_content = ""
    
    if main_content_element:
        # Process headings and paragraphs
        for element in main_content_element.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol', 'blockquote']):
            if element.name.startswith('h'):
                # Add spacing before headings (except the first one)
                if formatted_content:
                    formatted_content += "\n\n"
                
                # Format heading based on level
                heading_text = element.get_text(strip=True)
                if heading_text:
                    formatted_content += f"{heading_text}\n"
                    # Add underline for h1 and h2
                    if element.name in ['h1', 'h2']:
                        formatted_content += "=" * len(heading_text) + "\n"
            
            elif element.name == 'p':
                paragraph_text = element.get_text(strip=True)
                if paragraph_text:
                    if formatted_content:
                        formatted_content += "\n\n"
                    formatted_content += paragraph_text
            
            elif element.name == 'blockquote':
                quote_text = element.get_text(strip=True)
                if quote_text:
                    if formatted_content:
                        formatted_content += "\n\n"
                    formatted_content += "> " + quote_text
            
            elif element.name in ['ul', 'ol']:
                if formatted_content:
                    formatted_content += "\n\n"
                
                for i, li in enumerate(element.find_all('li', recursive=False)):
                    item_text = li.get_text(strip=True)
                    if item_text:
                        prefix = "• " if element.name == 'ul' else f"{i+1}. "
                        formatted_content += f"{prefix}{item_text}\n"
    
    # If still no structured content found, fall back to paragraphs with reasonable length
    if not formatted_content or len(formatted_content) < 200:
        paragraphs = soup.find_all('p')
        content_paragraphs = []
        
        for p in paragraphs:
            p_text = p.get_text(strip=True)
            # Consider paragraphs with reasonable length as content
            if len(p_text) > 40:
                content_paragraphs.append(p_text)
        
        if content_paragraphs:
            formatted_content = "\n\n".join(content_paragraphs)
    
    # If still no content, use structured extraction from the whole body
    if not formatted_content or len(formatted_content) < 100:
        body = soup.find('body')
        if body:
            # Extract headings and paragraphs from the entire body
            formatted_content = ""
            for element in body.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p']):
                if element.name.startswith('h'):
                    heading_text = element.get_text(strip=True)
                    if heading_text and len(heading_text) < 100:  # Avoid false headings
                        if formatted_content:
                            formatted_content += "\n\n"
                        formatted_content += f"{heading_text}\n"
                        if element.name in ['h1', 'h2']:
                            formatted_content += "=" * len(heading_text) + "\n"
                
                elif element.name == 'p':
                    p_text = element.get_text(strip=True)
                    if p_text and len(p_text) > 30:  # Skip very short paragraphs
                        if formatted_content:
                            formatted_content += "\n\n"
                        formatted_content += p_text
    
    # Clean the text
    formatted_content = clean_text(formatted_content)
    
    # Normalize whitespace - replace multiple newlines with double newlines
    formatted_content = re.sub(r'\n{3,}', '\n\n', formatted_content)
    
    return formatted_content

def download_file(url, save_dir):
    """Download a file and return its local path."""
    os.makedirs(save_dir, exist_ok=True)
    local_filename = os.path.join(save_dir, os.path.basename(url))
    with requests.get(url, stream=True) as r:
        r.raise_for_status()
        with open(local_filename, 'wb') as f:
            for chunk in r.iter_content(chunk_size=8192):
                f.write(chunk)
    return local_filename

def convert_pdf_to_text(pdf_file):
    """Convert PDF file to text with preserved formatting."""
    text = ""
    try:
        reader = PdfReader(pdf_file)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                # Normalize whitespace within the page
                page_text = re.sub(r' +', ' ', page_text)
                # Add page break between pages
                if text:
                    text += "\n\n"
                text += page_text
        
        # Normalize paragraph breaks
        text = re.sub(r'\n{3,}', '\n\n', text)
    except Exception as e:
        print(f"[ERROR] Failed to extract text from PDF: {e}")
    return text

def convert_docx_to_text(docx_file):
    """Convert DOCX file to text with preserved formatting."""
    text = ""
    try:
        doc = docx.Document(docx_file)
        
        # Process paragraphs
        for para in doc.paragraphs:
            para_text = para.text.strip()
            if para_text:
                # Check if it's a heading
                if para.style.name.startswith('Heading'):
                    # Add more spacing before headings
                    if text:
                        text += "\n\n"
                    text += para_text + "\n"
                else:
                    # Regular paragraph
                    if text:
                        text += "\n\n"
                    text += para_text
        
        # Process tables
        for table in doc.tables:
            if text:
                text += "\n\n"
            text += "Table:\n"
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text += row_text + "\n"
        
        # Normalize paragraph breaks
        text = re.sub(r'\n{3,}', '\n\n', text)
    except Exception as e:
        print(f"[ERROR] Failed to extract text from DOCX: {e}")
    return text

def get_domain_name(url):
    """Extract domain name from URL for folder naming."""
    parsed_url = urlparse(url)
    domain = parsed_url.netloc
    return domain.replace('.', '_')

def url_to_filename(url):
    """Convert URL to a safe filename using underscores."""
    # Remove protocol (http://, https://)
    filename = re.sub(r'^https?://', '', url)
    # Replace slashes and other problematic characters with underscores
    filename = re.sub(r'[/\\?%*:|"<>]', '_', filename)
    # Remove trailing underscores
    filename = filename.rstrip('_')
    # Ensure the filename isn't too long
    if len(filename) > 100:
        filename = filename[:100]
    return filename

def discover_urls(base_url, max_depth=1, limit_urls_control=25):
    """Discover all URLs on the site up to a given depth, one level at a time with URL limit control."""
    all_urls = []
    visited = set()
    
    # Start with the base URL at depth 0
    current_level = [(base_url.rstrip("/"), 0)]
    
    for depth in range(max_depth + 1):
        print(f"\nDiscovering URLs at depth {depth}...")
        next_level = []
        
        for url, url_depth in current_level:
            if url_depth != depth:
                continue
                
            normalized_url = url.rstrip("/")
            if normalized_url in visited:
                continue
                
            visited.add(normalized_url)
            all_urls.append((url, depth))
            
            # Check if we've reached the URL limit
            if len(all_urls) >= limit_urls_control:
                print(f"\n[LIMIT REACHED] Discovered {len(all_urls)} URLs, reached limit of {limit_urls_control}.")
                return all_urls
            
            try:
                print(f"Checking: {url}")
                resp = requests.get(url, timeout=10)
                if not resp.ok:
                    print(f"Failed to fetch {url}: HTTP {resp.status_code}")
                    continue
                    
                soup = BeautifulSoup(resp.text, "html.parser")
                
                # Collect links for the next depth level
                if depth < max_depth:
                    for link in soup.find_all("a", href=True):
                        href = link["href"]
                        full_url = urljoin(base_url, href)
                        full_url = full_url.split("#")[0].rstrip("/")
                        
                        if is_valid_url(full_url, base_url) and full_url not in visited:
                            next_level.append((full_url, depth + 1))
                            
            except Exception as e:
                print(f"[ERROR] Failed to fetch {url}: {e}")
        
        # Print summary for this depth
        depth_urls = [url for url, d in all_urls if d == depth]
        print(f"Found {len(depth_urls)} URLs at depth {depth}")
        
        # Move to the next level
        current_level = next_level
        
        # If we've completed the requested depth and found URLs, ask if user wants to continue
        if depth == max_depth and next_level and max_depth < 2:
            total_urls_found = len(all_urls)
            next_level_count = len(next_level)
            
            # Only ask if continuing won't exceed the limit immediately
            if total_urls_found + next_level_count <= limit_urls_control or limit_urls_control <= 0:
                print(f"\nFound {next_level_count} URLs at the next depth level.")
                user_input = input(f"Would you like to discover URLs at depth {depth+1} as well? (y/n): ")
                if user_input.lower() == 'y':
                    max_depth += 1
                
    return all_urls

def load_existing_hashes(hash_file_path):
    """Load existing hashes from file if it exists."""
    hashes = {}
    if os.path.exists(hash_file_path):
        with open(hash_file_path, 'r', encoding='utf-8') as f:
            for line in f:
                parts = line.strip().split(' ', 1)
                if len(parts) == 2:
                    hashes[parts[1]] = parts[0]
    return hashes

def scrape_and_save(main_url):
    """Main function to scrape website and save content."""
    # Step 1: Set up directories
    base_dir = "scraped_sites"
    os.makedirs(base_dir, exist_ok=True)
    domain_name = get_domain_name(main_url)
    site_dir = os.path.join(base_dir, domain_name)
    os.makedirs(site_dir, exist_ok=True)
    
    # Step 2: Get scraping parameters from user
    user_depth = input("Enter the depth level to check (1, 2, or 3): ")
    try:
        max_depth = int(user_depth)
        if max_depth < 1 or max_depth > 3:
            print("Invalid depth. Using default depth of 1.")
            max_depth = 1
    except ValueError:
        print("Invalid input. Using default depth of 1.")
        max_depth = 1
    
    # Get URL limit control parameter
    user_limit = input("Enter maximum number of URLs to discover (default is 25, enter 0 for no limit): ")
    try:
        limit_urls_control = int(user_limit)
        if limit_urls_control < 0:
            print("Invalid limit. Using default limit of 25 URLs.")
            limit_urls_control = 25
    except ValueError:
        print("Invalid input. Using default limit of 25 URLs.")
        limit_urls_control = 25
    
    # Step 3: Discover URLs up to the specified depth with the limit control
    print(f"Discovering URLs from {main_url} up to depth {max_depth} with limit of {limit_urls_control if limit_urls_control > 0 else 'unlimited'} URLs...")
    all_urls = discover_urls(main_url, max_depth=max_depth, limit_urls_control=limit_urls_control)
    
    # Group URLs by depth for display
    urls_by_depth = {}
    for url, depth in all_urls:
        if depth not in urls_by_depth:
            urls_by_depth[depth] = []
        urls_by_depth[depth].append(url)
    
    # Print summary of discovered URLs
    total_urls = len(all_urls)
    print("\nSummary of discovered URLs:")
    for depth in sorted(urls_by_depth.keys()):
        print(f"Depth {depth}: {len(urls_by_depth[depth])} URLs")
    
    print(f"\nTotal URLs discovered: {total_urls}")
    
    # Step 4: Ask user how many URLs to scrape
    user_input = input(f"How many URLs would you like to scrape? (1-{total_urls} or 'all'): ")
    
    if user_input.lower() == 'all':
        urls_to_scrape = all_urls.copy()
    else:
        try:
            num_to_scrape = int(user_input)
            if num_to_scrape <= 0:
                print("Invalid input. Scraping all URLs.")
                urls_to_scrape = all_urls.copy()
            elif num_to_scrape > total_urls:
                print(f"Only {total_urls} URLs available. Scraping all of them.")
                urls_to_scrape = all_urls.copy()
            else:
                # Prioritize by depth: main page first, then depth 1, then depth 2, etc.
                urls_to_scrape = []
                remaining = num_to_scrape
                
                for depth in sorted(urls_by_depth.keys()):
                    depth_urls = urls_by_depth[depth]
                    urls_to_add = depth_urls[:remaining]
                    urls_to_scrape.extend([(url, depth) for url in urls_to_add])
                    remaining -= len(urls_to_add)
                    if remaining <= 0:
                        break
        except ValueError:
            print("Invalid input. Scraping all URLs.")
            urls_to_scrape = all_urls.copy()
    
    # Step 5: Load existing hashes
    hash_file_path = os.path.join(site_dir, 'content_hashes.txt')
    existing_hashes = load_existing_hashes(hash_file_path)
    
    # Step 6: Scrape content level by level
    scraped_count = 0
    new_or_updated_count = 0
    
    for depth in sorted(urls_by_depth.keys()):
        # Filter urls_to_scrape to get only those at current depth
        depth_urls = []
        for url_info in urls_to_scrape:
            url, d = url_info
            if d == depth:
                depth_urls.append(url)
                
        print(f"\nScraping {len(depth_urls)} URLs at depth {depth}...")
        
        for i, url in enumerate(depth_urls):
            try:
                print(f"[{i+1}/{len(depth_urls)}] Scraping: {url}")
                resp = requests.get(url, timeout=10)
                if not resp.ok:
                    print(f"Failed to fetch {url}: HTTP {resp.status_code}")
                    continue
                
                # Process based on content type
                content_type = resp.headers.get('Content-Type', '').lower()
                
                if 'text/html' in content_type:
                    soup = BeautifulSoup(resp.text, "html.parser")
                    content = extract_main_content(soup)
                    
                elif any(url.lower().endswith(ext) for ext in TEXT_EXTENSIONS):
                    # Download and process document files
                    try:
                        file_name = download_file(url, site_dir)
                        
                        if file_name.lower().endswith(".pdf"):
                            content = convert_pdf_to_text(file_name)
                        elif file_name.lower().endswith(".docx"):
                            content = convert_docx_to_text(file_name)
                        else:  # .txt files
                            with open(file_name, 'r', encoding="utf-8", errors="replace") as f:
                                content = f.read()
                    except Exception as e:
                        print(f"[ERROR] Failed to process file {url}: {e}")
                        continue
                else:
                    print(f"Skipping unsupported content type: {content_type}")
                    continue
                
                # Skip if content is too short (likely not meaningful)
                if len(content) < 100:
                    print(f"Skipping {url}: Content too short ({len(content)} chars)")
                    continue
                
                # Compute hash of the content
                content_hash = compute_hash(content)
                
                # Check if content is new or updated
                if url in existing_hashes and existing_hashes[url] == content_hash:
                    print(f"Content unchanged for {url}, skipping")
                    scraped_count += 1
                    continue
                
                # Create a filename from the URL
                filename = f"depth_{depth}_{url_to_filename(url)}.txt"
                file_path = os.path.join(site_dir, filename)
                
                # Ensure directory exists before saving
                os.makedirs(os.path.dirname(file_path), exist_ok=True)
                
                # Save the content
                with open(file_path, "w", encoding="utf-8") as f:
                    f.write(f"URL: {url}\n\n{content}")
                
                # Update hash record
                existing_hashes[url] = content_hash
                
                # Save hash after each successful scrape
                with open(hash_file_path, 'w', encoding='utf-8') as f:
                    for url_key, hash_val in existing_hashes.items():
                        f.write(f"{hash_val} {url_key}\n")
                
                scraped_count += 1
                new_or_updated_count += 1
                
                # Small delay to be nice to the server
                time.sleep(random.uniform(0.5, 1.5))
                
            except Exception as e:
                print(f"[ERROR] Failed to scrape {url}: {e}")

    print(f"\nScraping completed!")
    print(f"Total URLs processed: {scraped_count}")
    print(f"New or updated content: {new_or_updated_count}")
    print(f"Results saved to: {site_dir}")
    return domain_name
